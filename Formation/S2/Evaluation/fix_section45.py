import docx
from docx.oxml.ns import qn
from copy import deepcopy

doc = docx.Document("/mnt/c/Users/shr3dr/Documents/Vaults/CRMEF/E-Portfolio/Formation/S2/Evaluation/Rapport du module evaluation - Soufiane Fhaili.docx")
body = doc.element.body
children = list(body.iterchildren())

def get_text(child):
    if child.tag != qn('w:p'):
        return None
    text = ''
    for t in child.iterchildren(qn('w:r')):
        tv = t.find(qn('w:t'))
        if tv is not None and tv.text:
            text += tv.text
    return text

# ===== FIX 1: Reorder 4.5 section content =====
# Find the reversed 4.5 section: from "R"flexion personnelle : Avant ce module" to heading
start_idx = None
end_idx = None
for i, child in enumerate(children):
    text = get_text(child)
    if text and "Avant ce module" in text and "communication" in text.lower().replace("\u2019", "'"):
        start_idx = i
    if text and "4.5 Communication" in text:
        end_idx = i
        break

if start_idx and end_idx and start_idx < end_idx:
    # Children start_idx to end_idx need to be reversed
    to_reverse = children[start_idx:end_idx+1]
    print(f"Reversing {len(to_reverse)} elements (children {start_idx}-{end_idx})")
    
    # Remove them
    for elem in to_reverse:
        body.remove(elem)
    
    # Reverse the order
    to_reverse.reverse()
    
    # Reinsert at the original start position (ref = child at start_idx or blank before 5.0)
    children = list(body.iterchildren())
    
    # Find the element that was right after the 4.4 section (child at start_idx before removal)
    # After removal, find "R"flexion personnelle : Ce qui me semble" (end of 4.4)
    ref_idx = None
    for i, child in enumerate(children):
        text = get_text(child)
        if text and "Ce qui me semble le plus important" in text:
            ref_idx = i
            break
    
    if ref_idx:
        ref = children[ref_idx]
        for elem in reversed(to_reverse):
            ref.addnext(elem)
        print(f"Reinserted reversed elements after child {ref_idx}")
    else:
        print("ERROR: Could not find insertion point")
else:
    print(f"ERROR: Section 4.5 not found properly. start={start_idx}, end={end_idx}")

# ===== FIX 2: Fix duplicate "mode, mode" in 4.4 =====
children = list(body.iterchildren())
for child in children:
    text = get_text(child)
    if text and "mode, mode" in text and "F et D" in text:
        for r in child.iterchildren(qn('w:r')):
            tv = r.find(qn('w:t'))
            if tv is not None and tv.text and "mode, mode" in tv.text:
                tv.text = tv.text.replace("mode, mode", "mode")
                print(f"Fixed duplicate 'mode, mode' -> 'mode'")
                break

# Also check for "mode, mode, F" in section 6
for child in children:
    text = get_text(child)
    if text and "Analyser syst" in text and "mode, mode" in text:
        for r in child.iterchildren(qn('w:r')):
            tv = r.find(qn('w:t'))
            if tv is not None and tv.text and "mode, mode" in tv.text:
                tv.text = tv.text.replace("mode, mode", "mode")
                print(f"Fixed duplicate in section 6")

output_path = "/mnt/c/Users/shr3dr/Documents/Vaults/CRMEF/E-Portfolio/Formation/S2/Evaluation/Rapport du module evaluation - Soufiane Fhaili.docx"
doc.save(output_path)
print(f"Saved. Tables: {len(doc.tables)}, Paras: {sum(1 for p in doc.paragraphs if p.text.strip())}")
