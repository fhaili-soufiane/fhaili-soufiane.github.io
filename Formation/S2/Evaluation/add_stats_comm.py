import docx
from docx import Document
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from copy import deepcopy

SQ = "\u2019"
MUL = "\u00d7"

doc = Document("/mnt/c/Users/shr3dr/Documents/Vaults/CRMEF/E-Portfolio/Formation/S2/Evaluation/Rapport du module evaluation - Soufiane Fhaili.docx")
body = doc.element.body
children = list(body.iterchildren())

def get_text(child):
    if child.tag != qn('w:p'):
        return None
    text = ''
    for t in child.iterchildren(qn('w:r')):
        tv = t.find(qn('w:t'))
        if tv is not None and tv.text:
            text += tv.text
    return text

def esc(text):
    text = text.replace('&', '&amp;')
    text = text.replace('<', '&lt;')
    text = text.replace('>', '&gt;')
    text = text.replace('"', '&quot;')
    return text

def make_para(text, style='Normal'):
    pPr = '<w:pPr><w:pStyle w:val="Normal"/></w:pPr>'
    if style == 'List Bullet':
        pPr = '<w:pPr><w:pStyle w:val="ListBullet"/><w:ind w:left="720"/></w:pPr>'
    p = parse_xml(f'<w:p {nsdecls("w")}>{pPr}</w:p>')
    r = parse_xml(f'<w:r {nsdecls("w")}><w:rPr><w:sz w:val="22"/></w:rPr><w:t xml:space="preserve">{esc(text)}</w:t></w:r>')
    p.find(qn('w:pPr')).addnext(r)
    return p

def make_blank():
    return parse_xml(f'<w:p {nsdecls("w")}><w:pPr><w:spacing w:after="0" w:line="240" w:lineRule="auto"/></w:pPr></w:p>')

def make_heading(text):
    p = parse_xml(f'<w:p {nsdecls("w")}><w:pPr><w:pStyle w:val="Normal"/></w:pPr></w:p>')
    r = parse_xml(f'<w:r {nsdecls("w")}><w:rPr><w:b/><w:sz w:val="24"/></w:rPr><w:t xml:space="preserve">{esc(text)}</w:t></w:r>')
    p.find(qn('w:pPr')).addnext(r)
    return p

def make_reflection(prefix, text):
    p = parse_xml(f'<w:p {nsdecls("w")}><w:pPr><w:pStyle w:val="Normal"/></w:pPr></w:p>')
    r1 = parse_xml(f'<w:r {nsdecls("w")}><w:rPr><w:i/><w:color w:val="4472C4"/><w:sz w:val="22"/></w:rPr><w:t xml:space="preserve">{esc(prefix)}</w:t></w:r>')
    r2 = parse_xml(f'<w:r {nsdecls("w")}><w:rPr><w:sz w:val="22"/></w:rPr><w:t xml:space="preserve">{esc(text)}</w:t></w:r>')
    p.find(qn('w:pPr')).addnext(r1)
    r1.addnext(r2)
    return p

def make_cell(text, width=2250):
    cell = parse_xml(f'''<w:tc {nsdecls("w")}>
        <w:tcPr><w:tcW w:w="{width}" w:type="dxa"/></w:tcPr>
        <w:p><w:pPr><w:pStyle w:val="Normal"/></w:pPr>
            <w:r><w:rPr><w:sz w:val="18"/></w:rPr><w:t xml:space="preserve">{esc(text)}</w:t></w:r>
        </w:p>
    </w:tc>''')
    return cell

# Find stats table
def find_stats_table():
    for child in children:
        if child.tag == qn('w:tbl'):
            rows = list(child.iterchildren(qn('w:tr')))
            if len(rows) >= 2:
                first = ''
                for cell in rows[0].iterchildren(qn('w:tc')):
                    for tp in cell.iterchildren(qn('w:p')):
                        for r in tp.iterchildren(qn('w:r')):
                            tv = r.find(qn('w:t'))
                            if tv is not None and tv.text:
                                first += tv.text
                    if first:
                        break
                if 'Indicateur' in first:
                    return child
    return None

# ============================================================
# 1. FIX DUPLICATE MODE ROW + ADD MAX/MIN/VARIANCE
# ============================================================
tbl = find_stats_table()
if tbl:
    rows = list(tbl.iterchildren(qn('w:tr')))
    print(f"Stats table: {len(rows)} rows (1 header + {len(rows)-1} data)")
    
    # Find and remove duplicate Mode rows (keep only the first)
    mode_indices = []
    for ri in range(1, len(rows)):
        first_cell = ''
        for cell in rows[ri].iterchildren(qn('w:tc')):
            for tp in cell.iterchildren(qn('w:p')):
                for r in tp.iterchildren(qn('w:r')):
                    tv = r.find(qn('w:t'))
                    if tv is not None and tv.text:
                        first_cell += tv.text
            if first_cell:
                break
        if first_cell.strip() == 'Mode':
            mode_indices.append(ri)
    
    # Remove all but first Mode row (remove from higher index first)
    if len(mode_indices) > 1:
        for idx in reversed(mode_indices[1:]):
            tbl.remove(rows[idx])
        print(f"   Removed {len(mode_indices)-1} duplicate Mode row(s)")
    
    # Re-fetch rows
    rows = list(tbl.iterchildren(qn('w:tr')))
    
    # Find where to insert new rows (after Écart-type = last row)
    # Add: Maximum, Minimum, Variance
    new_data = [
        ("Maximum", "Note la plus {SQ}lev{SQ}e dans la s{SQ}rie", "Permet de situer le meilleur niveau atteint et de d{SQ}finir l{SQ}exigence maximale.", "Utile pour v{SQ}rifier si l{SQ}preuve {SQ}tait accessible au meilleur."),
        ("Minimum", "Note la plus basse dans la s{SQ}rie", "Indique le niveau le plus faible. {SQ}carter avec la moyenne pour mesurer l{SQ}cart.", "Alerte sur les {SQ}l{SQ}ves en grande difficult{SQ}."),
        ("Variance", "Moyenne des carr{SQ}s des {SQ}carts {SQ} la moyenne", "Mesure th{SQ}orique de dispersion. La racine carr{SQ}e donne l{SQ}{SQ}cart-type.", "Utilis{SQ}e pour calculer l{SQ}{SQ}cart-type. Peu intuitive seule."),
    ]
    
    for label, formula, lecture, ref in new_data:
        row = parse_xml(f'<w:tr {nsdecls("w")}><w:trPr/></w:tr>')
        row.append(make_cell(label))
        row.append(make_cell(formula.replace("{SQ}", SQ)))
        row.append(make_cell(lecture.replace("{SQ}", SQ)))
        row.append(make_cell(ref.replace("{SQ}", SQ)))
        rows[-1].addnext(row)
    
    print(f"   Added Maximum, Minimum, Variance rows")
    
    # Verify
    rows = list(tbl.iterchildren(qn('w:tr')))
    print(f"   Now {len(rows)} rows: {[get_text(r) or '<table>' for r in [list(r.iterchildren(qn('w:tc')))[0] for r in rows[1:]]]}")
else:
    print("ERROR: Stats table not found")

# Update children
children = list(body.iterchildren())

# ============================================================
# 2. ADD SECTION 4.5: COMMUNICATION DES R"SULTATS
# ============================================================
# Find where to insert: after 4.4 section end (before 5.0)
# Find "5. Strat"gies" heading
idx_5 = None
for i, child in enumerate(children):
    text = get_text(child)
    if text and text.strip() == "5. Strat{SQ}gies de soutien et de rem{SQ}diation".replace(SQ, "'") or (text and text.strip().startswith("5. ") and "Strat" in text):
        idx_5 = i
        break

if idx_5 is None:
    print("ERROR: Section 5 heading not found")
else:
    sec45_heading = make_heading(f"4.5 Communication des r{SQ}sultats et bulletin des notes")
    sec45_content = [
        make_blank(),
        sec45_heading,
        make_para(f"L{SQ}{SQ}valuation ne prend tout son sens que lorsque ses r{SQ}sultats sont communiqu{SQ}s de mani{SQ}re claire, transparente et adapt{SQ}e {SQ} chaque interlocuteur. Le syllabus du module et la note minist{SQ}rielle n{SQ}184 insistent sur la n{SQ}cessit{SQ} de communiquer les r{SQ}sultats aux trois cat{SQ}gories d{SQ}acteurs concern{SQ}es."),
        make_blank(),
        make_para(f"1. Aux apprenants :"),
        make_para(f"Notes chiffr{SQ}es : chaque devoir donne lieu {SQ} une note sur 20, accompagn{SQ}e du bar{SQ}me d{SQ}taill{SQ} et des crit{SQ}res de correction.", "List Bullet"),
        make_para(f"Annotations : des commentaires pr{SQ}cis sur les erreurs commises et les pistes d{SQ}am{SQ}lioration, directement sur la copie.", "List Bullet"),
        make_para(f"Appr{SQ}ciations qualitatives : un bilan global qui valorise les progr{SQ}s et identifie les axes de travail, au-del{SQ} du simple chiffre.", "List Bullet"),
        make_blank(),
        make_para(f"2. {SQ} l{SQ}administration :"),
        make_para(f"Les notes sont transmises {SQ} l{SQ}administration via le bulletin des notes, dans un d{SQ}lai maximal de 15 jours apr{SQ}s le devoir (note 184).", "List Bullet"),
        make_para(f"Le rapport d{SQ}analyse des r{SQ}sultats (section 5.6) accompagne les notes pour {SQ}clairer les d{SQ}cisions p{SQ}dagogiques et administratives.", "List Bullet"),
        make_blank(),
        make_para(f"3. Aux parents et tuteurs :"),
        make_para(f"Le bulletin des notes reste le support officiel de communication. Il pr{SQ}sente les notes par mati{SQ}re, les moyennes et les appr{SQ}ciations.", "List Bullet"),
        make_para(f"Une communication r{SQ}guli{SQ}re (r{SQ}unions parents-profs, carnets de correspondance) permet d{SQ}impliquer les familles dans le suivi scolaire.", "List Bullet"),
        make_blank(),
        make_para(f"Le bulletin des notes est un document officiel qui synth{SQ}tise les r{SQ}sultats de l{SQ}{SQ}l{SQ}ve sur une p{SQ}riode donn{SQ}e. En contr{SQ}le continu, il int{SQ}gre les notes des DS ressources, du DS d{SQ}int{SQ}gration et des activit{SQ}s formatives, calcul{SQ}es selon la formule Mc = (MNr{MUL}4 + Ni{MUL}4 + Na{MUL}2) / 10."),
        make_blank(),
        make_reflection("R{SQ}flexion personnelle : ", f"Avant ce module, je consid{SQ}rais la communication des r{SQ}sultats comme une simple formalit{SQ} administrative. J{SQ}ai compris qu{SQ}elle est un acte p{SQ}dagogique {SQ} part enti{SQ}re : un {SQ}l{SQ}ve qui comprend sa note et les commentaires qui l{SQ}accompagnent est plus apte {SQ} progresser. Au coll{SQ}ge Ibn Toumert, je veillerai {SQ} r{SQ}diger des annotations constructives et {SQ} organiser des moments d{SQ}{SQ}change individualis{SQ} avec les {SQ}l{SQ}ves apr{SQ}s chaque {SQ}valuation sommative."),
    ]
    
    ref = children[idx_5]
    for elem in reversed(sec45_content):
        ref.addprevious(elem)
    print(f"2. Added 4.5 before section 5 (child {idx_5})")

# Update children
children = list(body.iterchildren())

# ============================================================
# 3. UPDATE 5.6 RAPPORT D'ANALYSE
# ============================================================
# Find the misplaced content that was moved to 5.6 (about distinction entre DS ressources)
# After it, add a point about communication
for i, child in enumerate(children):
    text = get_text(child)
    if text and "distinction permet d'identifier pr" in text.lower().replace(SQ, "'") and "cibl" in text:
        # Insert after this paragraph
        comm_paras = [
            make_blank(),
            make_para(f"Enfin, le rapport doit inclure une section de communication : notes chiffr{SQ}es, annotations qualitatives, et un projet de bulletin des notes {SQ} destination de l{SQ}administration et des parents."),
        ]
        for elem in reversed(comm_paras):
            child.addnext(elem)
        print(f"3. Updated 5.6 with communication mention")
        break

# ============================================================
# 4. UPDATE 5.7 CHECKLIST
# ============================================================
# Find the checklist table and add communication items
for child in children:
    if child.tag == qn('w:tbl'):
        rows = list(child.iterchildren(qn('w:tr')))
        if rows:
            first = ''
            for cell in rows[0].iterchildren(qn('w:tc')):
                for tp in cell.iterchildren(qn('w:p')):
                    for r in tp.iterchildren(qn('w:r')):
                        tv = r.find(qn('w:t'))
                        if tv is not None and tv.text:
                            first += tv.text
                if first:
                    break
            if first.strip() == 'Moment' or 'Moment' in first:
                # This is the checklist table (section 5.7)
                # Add rows for communication
                new_checklist_rows = [
                    ("Communication", "Notes transmises dans les 15 jours ? Annotations pr{SQ}sentes ? Bulletin pr{SQ}par{SQ} ?"),
                    ("Parents", "Les familles ont-elles {SQ}t{SQ} inform{SQ}es des r{SQ}sultats et des d{SQ}marches de rem{SQ}diation ?"),
                ]
                # Find last row
                chk_rows = list(child.iterchildren(qn('w:tr')))
                last_tr = chk_rows[-1]
                for label, content in new_checklist_rows:
                    new_tr = parse_xml(f'<w:tr {nsdecls("w")}><w:trPr/></w:tr>')
                    w = 4500
                    new_tr.append(make_cell(label.replace("{SQ}", SQ), width=w))
                    new_tr.append(make_cell(content.replace("{SQ}", SQ), width=w))
                    last_tr.addnext(new_tr)
                print(f"4. Updated 5.7 checklist with communication items")
                break

# ============================================================
# 5. UPDATE 6 MISE EN OEUVRE
# ============================================================
# Find the engagement about "Analyser syst"matiquement les r"sultats"
# and add an engagement about communication
for child in children:
    text = get_text(child)
    if text and "Analyser syst" in text and "matiquement les r" in text:
        # This is a bullet item, find the item after it and insert before
        # Or just insert a new bullet after the entire list
        # Find the non-bullet paragraph after the bullets (conclusion paragraph)
        pass

# Alternative: find the list items in section 6 and append after "Tenir un journal de bord"
for i, child in enumerate(children):
    text = get_text(child)
    if text and "Tenir un journal de bord" in text:
        new_items = [
            make_blank(),
            make_para(f"Communiquer les r{SQ}sultats de mani{SQ}re transparente : notes, annotations, appr{SQ}ciations qualitatives, dans les d{SQ}lais r{SQ}glementaires.", "List Bullet"),
            make_para(f"Pr{SQ}parer le bulletin des notes pour chaque semestre, conform{SQ}ment aux exigences de l{SQ}administration et aux besoins d{SQ}information des familles.", "List Bullet"),
        ]
        for elem in reversed(new_items):
            child.addnext(elem)
        print(f"5. Updated section 6 with communication engagements")
        break

# ============================================================
# SAVE
# ============================================================
output_path = "/mnt/c/Users/shr3dr/Documents/Vaults/CRMEF/E-Portfolio/Formation/S2/Evaluation/Rapport du module evaluation - Soufiane Fhaili.docx"
doc.save(output_path)
print(f"\nSaved! Tables: {len(doc.tables)}, Non-empty paras: {sum(1 for p in doc.paragraphs if p.text.strip())}")
