#!/usr/bin/env python3
"""
One-pass script:
1. Copies original report
2. For each week:
   - "Activités réalisées" → "Contexte" + rewritten paragraph
   - Removes "Observations pédagogiques" heading + paragraphs (non-animation only)
   - Inserts analysis table (Dimension | Observations | Pistes) BEFORE "Analyse et réflexion"
     (non-animation) or BEFORE "Bilan" (animation)
   - For animation weeks: adds "Analyse et réflexion" heading + paragraph after the table

Key: uses only text modifications + element removals + element insertions.
NEVER moves existing elements (avoids XML tree corruption).
"""

import shutil
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

SRC = "MSP/Rapport_de_Stage_Soufiane_Fhaili.docx"
DST = "MSP/Rapport_de_Stage_Soufiane_Fhaili_Final.docx"

CONTEXTES = {
    "2.1.": "Première journée de stage au Collège Ibn Toumert. Accompagné de ma tutrice Mme Amina EL MOUSAID, j'ai découvert le fonctionnement de la salle d'informatique et observé une séance d'introduction à l'environnement du tableur Excel auprès de deux groupes de 2ème année.",
    "2.2.": "Deuxième journée d'observation de Mme Amina EL MOUSAID sur le même thème de l'environnement du tableur Excel, dispensé aux deux autres groupes de 2ème année collège.",
    "2.3.": "Première observation de stagiaires pairs. Deux de mes collègues ont animé des séances sur les formules dans Excel : Saleh AIT TATA (08:30–10:30) et Oussama ETTAGNATI (10:30–12:30).",
    "2.4.": "Journée charnière : j'ai observé Youssef OUADRAHMANE sur les formules simples (08:30–10:30), puis j'ai animé ma toute première séance d'enseignement (A1) sur le même thème (10:30–12:30).",
    "2.5.": "Observation de deux stagiaires sur les fonctions SOMME, MOYENNE, MAX et MIN dans Excel : Oussama ETTAGNATI (08:30–10:30) et Saleh AIT TATA (10:30–12:30).",
    "2.6.": "Journée comprenant l'observation de Youssef OUADRAHMANE sur les fonctions (08:30–10:30) et ma deuxième prise en charge (A2) sur les fonctions MIN, MAX, SOMME et MOYENNE (10:30–11:30).",
    "2.7.": "Observation de Saleh AIT TATA qui a animé une séance de révision des formules et fonctions, suivi d'un cours sur la mise en forme des cellules et d'une évaluation formative.",
    "2.8.": "Retour après les vacances de printemps. Observation de la tutrice Mme Amina EL MOUSAID qui a organisé une évaluation pratique sur les fonctions et la mise en forme dans Excel.",
    "2.9.": "Observation de la tutrice sur une évaluation pratique des graphiques dans Excel, avec une introduction à la fonction SI.",
    "2.10.": "Observation de la tutrice lors d'une séance de révision sur les graphiques et la fonction SI, en préparation de ma dernière séance animée.",
    "2.11.": "Dernière journée de stage. Trois séances observées le matin (Ossama IBOURK, Youssef OUADRAHMANE, Saleh AIT TATA) suivies de ma troisième et dernière prise en charge (A3) sur l'insertion des graphiques (11:30–12:30).",
}

REFFLEXIONS = {
    "2.4.": "Cette première prise en charge m'a permis de mesurer l'écart entre la préparation théorique et la réalité de la classe. J'ai pris conscience de l'importance de la gestion du temps, de la nécessité de prévoir des exercices différenciés pour les élèves rapides, et de l'efficacité d'une situation problème contextualisée pour capter l'attention des collégiens.",
    "2.6.": "Ma deuxième séance animée m'a conforté dans ma progression : je me sens plus à l'aise qu'en A1. La contrainte de temps (1h au lieu de 2h) m'a obligé à être plus efficace dans mes transitions et à recentrer l'essentiel. L'exercice noté en fin de séance a maintenu l'attention des élèves.",
    "2.11.": "Cette dernière prise en charge a été la plus réussie des trois. Je ressentais une réelle aisance face à la classe. La situation problème sur le temps d'écran a immédiatement parlé aux élèves. L'exercice noté m'a permis de vérifier les acquis en temps réel. Je termine la séquence Tableur avec le sentiment du travail accompli et une progression tangible de A1 à A3.",
}

ANIMATION = {"2.4.", "2.6.", "2.11."}

# ── Analysis table data ──
ANALYSIS = {
    "2.1.": [
        ("Didactique",
         "La séance a porté sur la découverte de l'environnement du tableur Excel. L'enseignante a présenté l'interface : ruban, classeur, feuilles, cellules, lignes et colonnes. Les notions étaient adaptées au niveau des apprenants de 2AC. La progression était cohérente.",
         "Il serait utile de renforcer la structuration visuelle des notions de base, par exemple avec un schéma projeté ou une fiche récapitulative des éléments de l'interface."),
        ("Pédagogique",
         "La séance a combiné plusieurs démarches pédagogiques : questionnement, démonstration au vidéoprojecteur et exercice pratique. L'alternance entre collectif et individuel était équilibrée. La tutrice circule dans les rangs pour accompagner les élèves.",
         "Prévoir un dispositif de tutorat entre élèves pendant les TP afin d'accompagner les plus lents sans freiner les plus rapides."),
        ("Relationnel",
         "Le climat de classe était propice à l'apprentissage. L'enseignante a installé une relation de confiance par une communication bienveillante et une voix posée. Les élèves se sentent en sécurité pour poser des questions.",
         "Maintenir ce climat positif tout en encourageant davantage les élèves à expliciter leurs démarches à l'oral."),
    ],
    "2.2.": [
        ("Didactique",
         "Même thème que la semaine 1 (environnement du tableur) mais dispensé à des groupes différents. La tutrice a adapté son discours en raccourcissant certaines explications qui avaient bien fonctionné la première fois et en insistant davantage sur les points difficiles identifiés.",
         "Renforcer la structuration visuelle des notions de base, par exemple avec un schéma projeté ou une fiche récapitulative."),
        ("Pédagogique",
         "La séance a combiné plusieurs démarches pédagogiques : questionnement, démonstration et exercice pratique. Adaptation fine du discours pédagogique par rapport à la semaine 1.",
         "Prévoir un dispositif de tutorat entre élèves pendant les TP pour accompagner les plus lents."),
        ("Relationnel",
         "Le climat de classe était propice à l'apprentissage. L'enseignante a installé une relation de confiance. Élèves engagés et participants.",
         "Maintenir ce climat positif tout en encourageant davantage les élèves à expliciter leurs démarches."),
    ],
    "2.3.": [
        ("Didactique",
         "La séance était construite autour d'une situation concrète liée à l'achat de fournitures. La progression du rappel vers l'application pratique était logique. L'utilisation des adresses de cellules a bien été comprise par la majorité des élèves.",
         "Consolider la notion d'adresse absolue, encore difficile pour plusieurs élèves, par un exercice de repérage dédié."),
        ("Pédagogique",
         "La séance a alterné rappel, situation-problème, activité pratique et remédiation. L'enseignant stagiaire a suivi une démarche structurée avec un diaporama détaillé.",
         "Prévoir plus d'exercices guidés sur la recopie des formules et l'utilisation du signe « = ». Distribuer la trace écrite en fin de séance."),
        ("Relationnel",
         "La relation avec les élèves était positive. Les élèves ont participé au rappel et à la situation problème. Climat de classe agréable.",
         "Maintenir ce climat positif, encourager les élèves à expliquer leurs démarches et solliciter davantage les plus réservés."),
    ],
    "2.4.": [
        ("Didactique",
         "Situation problème contextualisée (librairie) efficace — les élèves se sont immédiatement projetés. Progression du concret vers l'abstrait (formules simples → adresses relatives). Alternance français/arabe pour faciliter la compréhension.",
         "Insister sur le signe = dès le rappel. Prévoir un exercice supplémentaire pour les plus rapides. Introduire l'adresse absolue plus tôt."),
        ("Pédagogique",
         "Timing globalement respecté (2h). Alternance de phases claires. Supports : vidéoprojecteur, fichiers Excel, fiche pédagogique. Gestion de l'espace : déplacement dans toute la salle.",
         "Chronométrer chaque phase. Préparer des exercices différenciés (niveaux 1 et 2). Distribuer la trace écrite plus tôt pour gagner du temps."),
        ("Relationnel",
         "Élèves engagés et intéressés — l'activité pratique motive. Un binôme bavardait : avertissement oral. Climat général positif.",
         "Anticiper les règles de travail en début de séance. Gérer les interruptions plus fermement dès le début."),
    ],
    "2.5.": [
        ("Didactique",
         "La séance s'est inscrite dans la continuité du travail sur les formules. L'introduction des fonctions SOMME, MOYENNE, MAX, MIN a été bien amenée avec une situation problème bibliothèque.",
         "Consolider l'explication de l'adresse absolue et du symbole « $ ». Proposer un exercice de repérage de la syntaxe correcte entre « ; » et « : »."),
        ("Pédagogique",
         "Organisation cohérente alternant rappel, mise en situation, démonstration, exercices et synthèse. L'enseignant stagiaire circule dans les rangs.",
         "Privilégier davantage la pratique individuelle. Prévoir plus de temps pour la manipulation autonome."),
        ("Relationnel",
         "Climat de classe globalement positif. Les élèves ont participé activement. Bonne interaction entre l'enseignant et les apprenants.",
         "Maintenir ce climat de participation en sollicitant davantage les élèves réservés."),
    ],
    "2.6.": [
        ("Didactique",
         "Situation problème contextualisée (bibliothèque) efficace — chaque fonction a trouvé son utilité. Progression respectée. Lien entre les quatre fonctions bien établi.",
         "Introduire les fonctions une par une avec plus de temps de manipulation. Prévoir un exercice de repérage de la syntaxe correcte."),
        ("Pédagogique",
         "Timing respecté malgré la contrainte de 1h. Alternance de 7 phases courtes. Supports : vidéoprojecteur, fichiers Excel. Exercice noté en fin de séance.",
         "Chronométrer chaque phase avec une marge. Préparer une version allégée de la trace écrite."),
        ("Relationnel",
         "Bonne interactivité. Élèves actifs et impliqués. L'exercice noté a maintenu l'attention. Climat serein.",
         "Maintenir cette dynamique positive. Varier les sollicitations pour inclure tous les élèves."),
    ],
    "2.7.": [
        ("Didactique",
         "La séance a suivi une progression cohérente, enchaînant logiquement rappels, exercices de révision, démonstration de mise en forme et évaluation formative.",
         "Élaborer une fiche récapitulative des fonctions et de la mise en forme. Prévoir un exercice de synthèse."),
        ("Pédagogique",
         "L'enseignant stagiaire a assuré un accompagnement régulier des élèves. Évaluation formative intégrée pertinente.",
         "Mettre en place des exercices différenciés. Chronométrer chaque phase plus rigoureusement."),
        ("Relationnel",
         "Les élèves se sont montrés engagés et intéressés. Climat de classe positif et propice aux apprentissages.",
         "Encourager davantage les élèves les plus réservés à participer et valoriser leurs contributions."),
    ],
    "2.8.": [
        ("Didactique",
         "Les objectifs de l'évaluation étaient en adéquation avec les apprentissages précédents. Consignes claires, critères d'évaluation explicités.",
         "Diversifier les situations d'évaluation. Proposer des exercices de remédiation ciblés."),
        ("Pédagogique",
         "Bonne gestion du temps et de l'espace. L'enseignante a circulé pour accompagner les élèves en difficulté.",
         "Renforcer l'accompagnement individualisé et prévoir davantage de temps pour la correction collective."),
        ("Relationnel",
         "Climat de classe positif et respectueux. Communication claire avec les apprenants.",
         "Encourager les interactions entre apprenants et valoriser systématiquement les efforts."),
    ],
    "2.9.": [
        ("Didactique",
         "La séance structurée autour d'Excel avec une évaluation pratique des graphiques et une introduction à la fonction SI. Approche interdisciplinaire mobilisée.",
         "Prévoir une trace écrite courte à la fin de chaque étape et afficher les critères d'évaluation."),
        ("Pédagogique",
         "Bien organisée avec une approche par compétences. Alternance entre rappel, TP, correction et évaluation.",
         "Mieux équilibrer le temps entre les phases. Laisser plus de temps à la manipulation autonome."),
        ("Relationnel",
         "Relation positive avec les élèves. L'enseignante encourageait les apprenants. Climat serein.",
         "Continuer à encourager les élèves et solliciter davantage les plus réservés."),
    ],
    "2.10.": [
        ("Didactique",
         "Graphiques : bon rappel, question stimulante concrète. Fonction SI : bonne introduction. Progression adaptée.",
         "Prévoir un exercice SI à 2 conditions pour les rapides. Créer une fiche récapitulative des fonctions."),
        ("Pédagogique",
         "Respect du temps global. Alternance équilibrée. Accompagnement individuel dans les rangs.",
         "Chronométrer chaque phase plus rigoureusement. Préparer un fichier Excel vierge prêt à l'emploi."),
        ("Relationnel",
         "Climat serein et respectueux. Élèves engagés et motivés. Relation positive.",
         "Utiliser un signal non verbal pour recentrer l'attention. Valoriser les efforts."),
    ],
    "2.11.": [
        ("Didactique",
         "Situation problème attractive (temps d'écran) — les élèves se sont reconnus. Lien interdisciplinaire Histoire-Géographie bien exploité.",
         "Proposer un exercice de « diagnostic » où les élèves associent un type de graphique à un jeu de données."),
        ("Pédagogique",
         "Timing de 1h respecté. Alternance de phases claires. Supports variés. Exercice d'évaluation noté intégré.",
         "Laisser plus de temps pour l'exploration autonome des types de graphiques."),
        ("Relationnel",
         "Élèves engagés et motivés. Situation problème a créé un intérêt immédiat. Aucun incident. Climat serein.",
         "Maintenir cette approche contextualisée. Continuer à lier les notions aux autres matières."),
    ],
}

MARKERS = ["2.1.", "2.2.", "2.3.", "2.4.", "2.5.", "2.6.", "2.7.", "2.8.", "2.9.", "2.10.", "2.11."]


# ═══════════════════════════════════════════════════
#  XML helpers
# ═══════════════════════════════════════════════════

def el(tag):
    """Create element in the w: namespace."""
    return OxmlElement(tag)


def make_borders():
    b = el('w:tblBorders')
    for name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        brd = el(f'w:{name}')
        brd.set(qn('w:val'), 'single')
        brd.set(qn('w:sz'), '4')
        brd.set(qn('w:space'), '0')
        brd.set(qn('w:color'), 'BFBFBF')
        b.append(brd)
    return b


def make_table_xml(rows_data):
    """Build a <w:tbl> element for the 3-column analysis table."""
    headers = ["Dimension", "Observations", "Pistes d'amélioration"]

    tbl = el('w:tbl')

    # tblPr
    tblPr = el('w:tblPr')
    tw = el('w:tblW')
    tw.set(qn('w:w'), '9000')
    tw.set(qn('w:type'), 'dxa')
    tblPr.append(tw)
    tblPr.append(make_borders())
    tblPr.append(el('w:tblLook'))
    tbl.append(tblPr)

    # tblGrid
    tg = el('w:tblGrid')
    for w_val in ['1200', '3900', '3900']:
        gc = el('w:gridCol')
        gc.set(qn('w:w'), w_val)
        tg.append(gc)
    tbl.append(tg)

    def make_cell(text, is_header=False, is_first_col=False):
        tc = el('w:tc')
        tcPr = el('w:tcPr')
        # shading
        shd = el('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        if is_header:
            shd.set(qn('w:fill'), '595959')
        elif is_first_col:
            shd.set(qn('w:fill'), 'F2F2F2')
        tcPr.append(shd)
        # width
        tcW = el('w:tcW')
        tcW.set(qn('w:w'), '1200' if is_first_col else '3900')
        tcW.set(qn('w:type'), 'dxa')
        tcPr.append(tcW)
        tc.append(tcPr)
        # para
        p = el('w:p')
        pPr = el('w:pPr')
        rPr = el('w:rPr')
        if is_header or is_first_col:
            b = el('w:b')
            rPr.append(b)
        if is_header:
            c = el('w:color')
            c.set(qn('w:val'), 'FFFFFF')
            rPr.append(c)
        sz = el('w:sz')
        sz.set(qn('w:val'), '20')
        rPr.append(sz)
        pPr.append(rPr)
        p.append(pPr)
        r = el('w:r')
        trPr = el('w:rPr')
        if is_header or is_first_col:
            b2 = el('w:b')
            trPr.append(b2)
        if is_header:
            c2 = el('w:color')
            c2.set(qn('w:val'), 'FFFFFF')
            trPr.append(c2)
        sz2 = el('w:sz')
        sz2.set(qn('w:val'), '20')
        trPr.append(sz2)
        r.append(trPr)
        t = el('w:t')
        t.text = text
        t.set(qn('xml:space'), 'preserve')
        r.append(t)
        p.append(r)
        tc.append(p)
        return tc

    # Header row
    tr_h = el('w:tr')
    for i, htext in enumerate(headers):
        tr_h.append(make_cell(htext, is_header=True, is_first_col=(i == 0)))
    tbl.append(tr_h)

    # Data rows
    for dim, obs, pistes in rows_data:
        tr = el('w:tr')
        tr.append(make_cell(dim, is_first_col=True))
        tr.append(make_cell(obs))
        tr.append(make_cell(pistes))
        tbl.append(tr)

    return tbl


def make_heading3(text):
    """Create a Heading 3 paragraph element."""
    p = el('w:p')
    pPr = el('w:pPr')
    ps = el('w:pStyle')
    ps.set(qn('w:val'), 'Heading3')
    pPr.append(ps)
    p.append(pPr)
    r = el('w:r')
    t = el('w:t')
    t.text = text
    t.set(qn('xml:space'), 'preserve')
    r.append(t)
    p.append(r)
    return p


def make_normal(text):
    p = el('w:p')
    r = el('w:r')
    t = el('w:t')
    t.text = text
    t.set(qn('xml:space'), 'preserve')
    r.append(t)
    p.append(r)
    return p


def make_empty():
    return el('w:p')


def get_para_text(elem):
    return ''.join(t.text or '' for t in elem.iter(qn('w:t'))).strip()


def get_heading_text(elem):
    return get_para_text(elem)


def is_heading(elem):
    pStyle = elem.find(qn('w:pPr') + '/' + qn('w:pStyle'))
    if pStyle is not None:
        sv = pStyle.get(qn('w:val'))
        return sv and ('Heading' in sv or 'Titre' in sv)
    return False


def get_style_id(elem):
    pStyle = elem.find(qn('w:pPr') + '/' + qn('w:pStyle'))
    if pStyle is not None:
        return pStyle.get(qn('w:val'))
    return None


def set_para_text(elem, new_text):
    """Set paragraph text for a heading element. Preserves first run of heading."""
    runs = list(elem.iter(qn('w:t')))
    for r in runs:
        r.text = ''
    if runs:
        runs[0].text = new_text


def remove_safe(elem):
    p = elem.getparent()
    if p is not None:
        p.remove(elem)


def insert_before(ref, new_elem):
    """Insert new_elem as the immediate previous sibling of ref."""
    ref.addprevious(new_elem)


# ═══════════════════════════════════════════════════
#  Main
# ═══════════════════════════════════════════════════

def main():
    shutil.copy2(SRC, DST)
    doc = Document(DST)
    body = doc.element.body
    print(f"Starting with {len(doc.paragraphs)} paragraphs, {len(doc.tables)} tables")

    for marker in MARKERS:
        is_anim = marker in ANIMATION
        context = CONTEXTES[marker]
        analysis_rows = ANALYSIS[marker]
        reflection = REFFLEXIONS.get(marker)

        # ── Find the H2 heading for this week ──
        h2_elem = None
        for child in body:
            if child.tag.endswith('}p') and is_heading(child):
                if get_heading_text(child).startswith(marker):
                    h2_elem = child
                    break
        if h2_elem is None:
            print(f"  ✗ {marker}: H2 not found")
            continue

        # ── Scan forward to identify key elements ──
        activites_h = None    # "Activités réalisées"
        activites_paras = []  # Normal paras under it
        obs_ped_h = None
        obs_ped_paras = []
        analyse_h = None      # "Analyse et réflexion" (non-animation only)
        analyse_para = None
        animes_h = None       # "Première séance animée..." (animation only)
        animes_paras = []
        animes_tables = []    # summary tables for animated sessions
        bilan_h = None        # "Bilan"
        other_pre_bilan = []  # everything else before Bilan

        phase = 'start'
        current = h2_elem.getnext()
        while current is not None:
            tag = current.tag.split('}')[-1]

            # Check for next section H2 → stop
            if tag == 'p' and is_heading(current) and current is not h2_elem:
                ht = get_heading_text(current)
                if ht.startswith("3.") or "Apports" in ht:
                    break
                if ht.startswith("2.") and not ht.startswith(marker):
                    break

            if tag == 'p':
                pt = get_para_text(current)
                ih = is_heading(current)

                if ih and pt == "Activités réalisées":
                    activites_h = current
                    phase = 'activites'
                elif ih and pt == "Observations pédagogiques":
                    obs_ped_h = current
                    phase = 'obs_ped'
                elif ih and pt == "Analyse et réflexion":
                    analyse_h = current
                    phase = 'analyse'
                elif ih and ("Première séance animée" in pt or "Deuxième séance animée" in pt or "Dernière séance animée" in pt):
                    animes_h = current
                    phase = 'animes'
                elif ih and pt == "Bilan":
                    bilan_h = current
                    break
                else:
                    if phase == 'activites':
                        activites_paras.append(current)
                    elif phase == 'obs_ped':
                        obs_ped_paras.append(current)
                    elif phase == 'analyse':
                        analyse_para = current
                    elif phase in ('start', 'animes', 'other'):
                        pass  # will handle these later

            elif tag == 'tbl':
                if phase == 'animes':
                    animes_tables.append(current)

            current = current.getnext()

        if bilan_h is None:
            print(f"  ✗ {marker}: Bilan not found")
            continue

        # ════════════════════════════════════════
        #  Apply modifications
        # ════════════════════════════════════════

        # 1. Activités réalisées → Contexte
        if activites_h is not None:
            set_para_text(activites_h, "Contexte")
            # Remove old paragraphs
            for p in activites_paras:
                remove_safe(p)
            # Insert context paragraph right after heading
            ctx_p = make_normal(context)
            activites_h.addnext(ctx_p)

        # 2. Remove Observations pédagogiques (non-animation)
        if obs_ped_h is not None:
            for p in obs_ped_paras:
                remove_safe(p)
            remove_safe(obs_ped_h)

        # 3. Determine insertion point for the analysis table
        # Non-animation: insert BEFORE "Analyse et réflexion"
        # Animation: insert BEFORE "Bilan" (after animated session content)
        if not is_anim and analyse_h is not None:
            insert_target = analyse_h
        else:
            insert_target = bilan_h

        # 4. Insert analysis table + heading
        if insert_target is not None:
            table_heading = make_heading3("Analyse dimensionnelle de la séance")
            blank = make_empty()
            table_elem = make_table_xml(analysis_rows)

            # Insert in order (each addprevious inserts immediately before target)
            # We want: heading → blank → table → target
            insert_before(insert_target, table_heading)  # ... [H3] [target]
            insert_before(insert_target, blank)           # ... [H3] [blank] [target]
            insert_before(insert_target, table_elem)      # ... [H3] [blank] [TBL] [target]

        # 5. For animation weeks: add "Analyse et réflexion" after table, before Bilan
        if is_anim and reflection is not None and insert_target is not None:
            refl_h = make_heading3("Analyse et réflexion")
            refl_p = make_normal(reflection)
            insert_before(insert_target, refl_h)  # heading first
            insert_before(insert_target, refl_p)  # then paragraph (below heading)

        print(f"  ✓ {marker} ({'anim' if is_anim else 'obs'})")

    doc.save(DST)
    print(f"\nSaved → {DST}")
    print(f"Final: {len(doc.paragraphs)} paragraphs, {len(doc.tables)} tables")


if __name__ == "__main__":
    main()
