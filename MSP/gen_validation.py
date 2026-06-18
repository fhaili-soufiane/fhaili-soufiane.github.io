from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

def set_cell_shading(cell, color):
    shading = OxmlElement('w:shd')
    shading.set(qn('w:val'), 'clear')
    shading.set(qn('w:color'), 'auto')
    shading.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading)

def add_table(doc, rows_data, header, col_widths=None):
    table = doc.add_table(rows=len(rows_data) + 1, cols=len(header))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(header):
        cell = table.rows[0].cells[j]
        cell.text = ''
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(10)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, '1E3A5F')
    for i, row_data in enumerate(rows_data):
        for j, val in enumerate(row_data):
            cell = table.rows[i + 1].cells[j]
            cell.text = ''
            run = cell.paragraphs[0].add_run(val)
            run.font.size = Pt(10)
    if col_widths:
        for row in table.rows:
            for j, w in enumerate(col_widths):
                row.cells[j].width = Cm(w)
    doc.add_paragraph('')
    return table

def add_heading(doc, text, level=13):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(level)
    run.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)

def add_body(doc, text):
    p = doc.add_paragraph(text)
    p.paragraph_format.first_line_indent = Cm(0.5)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = Pt(14)

# ═══════════════════════
# TITLE
# ═══════════════════════
add_heading(doc, "Fiche d'analyse APP — Séance animée", 18)
p = doc.add_paragraph()
r = p.add_run('Exemple : Les formules simples dans un tableur (26/02/2024)')
r.italic = True
r.font.size = Pt(11)
r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
doc.add_paragraph('')

# ── 1. Infos générales ──
add_heading(doc, '1. Informations générales', 13)
table = doc.add_table(rows=5, cols=4)
table.style = 'Table Grid'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
data = [
    ['Stagiaire', 'Soufiane FHAILI', 'Établissement', 'Collège Ibn Toumert'],
    ['Classe', '2AC', 'Effectif', '34 élèves'],
    ['Date', '26/02/2024', 'Horaire', '08h30 — 10h30'],
    ['Thème', 'Les formules simples dans un tableur', 'Type', 'Prise en charge'],
    ['Durée', '2h', 'Tuteur', 'Mme HRICH'],
]
for i, row_data in enumerate(data):
    for j, val in enumerate(row_data):
        cell = table.rows[i].cells[j]
        cell.text = ''
        run = cell.paragraphs[0].add_run(val)
        run.font.size = Pt(10)
        if j % 2 == 0:
            run.bold = True
            set_cell_shading(cell, 'E8EDF2')
doc.add_paragraph('')

# ── 2. Contexte ──
add_heading(doc, '2. Contexte de la séance (Description)', 13)
add_body(doc,
    "Cette séance s'inscrit dans la séquence Tableur (Unité 3). Lors de la séance précédente, "
    "les élèves ont découvert l'environnement Excel, l'insertion de données et les types de données "
    "(texte, valeur numérique, date).\n\n"
    "Objectif : les élèves seront capables de manipuler les formules simples (+, -, *, /) avec les "
    "adresses de cellules et de distinguer les adresses relatives des adresses absolues.\n\n"
    "Préparation : j'ai préparé une fiche pédagogique, un diaporama pour la situation problème, "
    "des fichiers Excel pour la phase de mise en commun et un exercice d'application pour le "
    "réinvestissement. C'était ma première prise en charge d'une séance complète de 2h."
)

# ── 3. Déroulement ──
add_heading(doc, '3. Déroulement chronologique', 13)
add_table(doc, [
    ['Mise en situation', '15 min',
     'Rappel des prérequis (environnement Excel, types de données). '
     'Projection de la situation problème (librairie de Karim). '
     'Questions guidées pour faire émerger le problème.',
     'Comprennent la situation. Ouvrent leur classeur. '
     'Créent une nouvelle feuille "TP5". Copient le tableau.'],
    ['Analyse', '10 min',
     'Discussion collective pour valider la compréhension. '
     'Demande : « Comment faire pour que Excel calcule le total des cahiers automatiquement ? » '
     'Note les hypothèses au tableau.',
     'Proposent des hypothèses : "taper 5*4.50", "taper =5*4.5". '
     'Un élève propose =5*4.5 — je valide.']
    ,
    ['Mise en commun', '25 min',
     'Démonstration : taper =B2*C2. Explication : « Excel utilise les adresses des cellules ». '
     'Montre la recopie de la formule dans D3:D5 (adresse relative). '
     'Test : changer le prix en B2 → la mise à jour est automatique.',
     'Testent la solution sur leur poste. Changent le prix du cahier. '
     'Certains confondent adresse relative et absolue. '
     'Je ré-explique avec le déplacement de la formule.'],
    ['Institutionnalisation', '15 min',
     'Trace écrite projetée : définition d\'une formule, syntaxe, '
     'adresses relatives. Distribution de la fiche récapitulative.',
     'Copient la trace. Posent des questions sur la syntaxe.'],
    ['Réinvestissement', '25 min',
     'Exercice : dans C6 insérer "prix total général", '
     'dans D6 calculer la somme avec =SOMME(D2:D5). '
     'Circule entre les rangs, aide individualisée.',
     'Travail individuel sur Excel. ~70% réussissent. '
     'Les plus rapides aident les voisins. '
     'Certains oublient le signe =, je rappelle.'],
    ['Correction', '20 min',
     'Correction collective au vidéoprojecteur. '
     'Questions de synthèse sur l\'utilité des formules.',
     'Comparent leur travail. Participent à la correction.'],
    ['Clôture', '10 min',
     'Synthèse orale : « Qu\'avons-nous appris aujourd\'hui ? » '
     'Annonce de la prochaine séance : les fonctions.',
     'Répondent : les formules, les adresses, SOMME. '
     'Objectif atteint.'],
], ['Phase', 'Durée', 'Activité enseignant', 'Activité élèves'],
   col_widths=[3.5, 1.5, 6, 6])

# ── 4. Problématisation ──
add_heading(doc, '4. Problématisation', 13)
add_body(doc,
    "Plusieurs questions émergent de cette séance :\n"
    "• Comment rendre intuitive la différence entre adresse relative et absolue ?\n"
    "• Quel équilibre entre cours magistral et pratique individuelle sur machine ?\n"
    "• Comment gérer l'hétérogénéité quand les élèves avancent à des rythmes différents ?\n"
    "• Le travail en binôme favorise-t-il l'apprentissage ou crée-t-il de la dépendance ?"
)

# ── 5. Analyse ──
add_heading(doc, '5. Analyse par dimensions', 13)
add_table(doc, [
    ['Didactique',
     'Situation de départ, construction, gestion des erreurs, évaluation',
     'Situation problème contextualisée (librairie) efficace — les élèves se sont immédiatement projetés. '
     'Progression : du concret (multiplication simple) vers l\'abstrait (adresses, recopie). '
     'Confusion relative/absolue gérée par re-explication avec analogie du déplacement. '
     '70% de réussite à l\'exercice de réinvestissement. Certains oublient le signe =.',
     'Insister sur le signe = dès le rappel. '
     'Prévoir un exercice supplémentaire pour les plus rapides.'],
    ['Pédagogique',
     'Gestion du temps, méthodes, supports',
     'Timing globalement respecté (2h). Alternance de 7 phases. '
     'Supports : vidéoprojecteur, fichiers Excel, fiche pédagogique, trace écrite. '
     'Salle disposée en U — bonne visibilité. '
     'La phase de réinvestissement a pris plus de temps que prévu.',
     'Chronométrer chaque phase. '
     'Préparer des exercices différenciés (niveaux 1 et 2).'],
    ['Relationnel',
     'Climat, communication, discipline',
     'Élèves engagés et intéressés — l\'activité pratique motive. '
     'Un binôme bavardait : avertissement oral. '
     'Un élève aidait ses voisins sans permission : rappel des consignes. '
     'Le tuteur est intervenu pour m\'aider à garder le calme. '
     'Climat général positif.',
     'Anticiper les règles de travail en début de séance. '
     'Gérer les interruptions plus fermement dès le début.'],
], ['Dimension', 'Éléments observés', 'Observations', "Pistes d'amélioration"],
   col_widths=[2.5, 4, 5.5, 4.5])

# ── 6. Théorisation ──
add_heading(doc, '6. Théorisation', 13)
add_body(doc,
    "Plusieurs concepts pédagogiques éclairent cette séance :\n"
    "• **Situation problème** (Meirieu) : la contextualisation dans une librairie a donné du sens "
    "aux apprentissages et motivé les élèves.\n"
    "• **Conflit cognitif** (Piaget) : le changement de prix qui ne mettait pas à jour le calcul "
    "a créé un déséquilibre, amenant les élèves à chercher la solution avec les adresses.\n"
    "• **Médiation / étayage** (Vygotsky) : le guidage pas à pas et l'analogie du déplacement "
    "ont servi de soutien pour franchir la difficulté.\n"
    "• **Zone proximale de développement** : l'exercice guidé (mise en commun) suivi de "
    "l'exercice autonome (réinvestissement) a respecté cette progression.\n"
    "• **Évaluation formative** : la circulation et la correction collective ont permis "
    "un feedback immédiat et une régulation."
)

# ── 7. Synthèse réflexive ──
add_heading(doc, '7. Synthèse réflexive personnelle', 13)
add_body(doc,
    "Pour une première prise en charge de 2h, je considère la séance réussie. J'ai réussi à "
    "structurer la séance en phases cohérentes, à capter l'attention des élèves et à gérer "
    "les interactions. La situation problème a été un levier efficace.\n\n"
    "Points à améliorer : consacrer plus de temps à la phase de réinvestissement, préparer "
    "des exercices différenciés pour gérer l'hétérogénéité, et chronométrer chaque phase "
    "plus rigoureusement.\n\n"
    "Ce que j'en retiens : l'importance de la gestion de la parole, de l'espace et de la "
    "bienveillance. L'alternance théorie-pratique est essentielle pour maintenir l'attention "
    "sur une plage de 2h. Pour la prochaine séance (Les fonctions), je préparerai un "
    "exercice de niveau 2 pour les élèves rapides et je sera plus strict sur les règles "
    "de travail en début de séance."
)

# ── Save ──
output = '/tmp/Fiche_APP_Seance_Tableur.docx'
doc.save(output)
print(f'✅ Généré : {output}')
