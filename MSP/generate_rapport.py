from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

doc = Document()

# ── Styles ──
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

# ── Helper ──
def set_cell_shading(cell, color):
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:val'), 'clear')
    shading_elm.set(qn('w:color'), 'auto')
    shading_elm.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading_elm)

def add_header_row(table, texts, color='1E3A5F'):
    row = table.rows[0]
    for i, text in enumerate(texts):
        cell = row.cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(text)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(10)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, color)

def add_row(table, cells_texts):
    row = table.add_row()
    for i, text in enumerate(cells_texts):
        cell = row.cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(text)
        run.font.size = Pt(10)

# ── Cover page ──
for _ in range(6):
    doc.add_paragraph('')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Rapport de Stage\nMise en Situations Professionnelles')
run.bold = True
run.font.size = Pt(26)
run.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)

doc.add_paragraph('')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('CRMEF Ibn Roch Marrakech\nFilière Collège — Spécialité Informatique\nAnnée de Formation 2025-2026')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

doc.add_paragraph('')
doc.add_paragraph('')

info = [
    ('Stagiaire :', 'Soufiane FH AILI'),
    ('Établissement :', 'Collège Ibn Toumert, Marrakech'),
    ('Niveau :', '2ème année collège'),
    ('Tuteur :', 'M. El Amrani'),
    ('Période :', 'Octobre 2025 — Juin 2026'),
    ('Soutenance :', '19 Juillet 2026'),
]
for label, value in info:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f'{label} ')
    run.bold = True
    run.font.size = Pt(12)
    run = p.add_run(value)
    run.font.size = Pt(12)

doc.add_page_break()

# ── Table des matieres ──
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Table des Matières')
run.bold = True
run.font.size = Pt(18)
run.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)

doc.add_paragraph('')
toc = [
    'Introduction',
    'Présentation de l\'établissement',
    'Fiches de séance (Séances 1 à 10)',
    'Synthèse et bilan',
    'Annexes',
]
for item in toc:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(4)
    run = p.runs[0]
    run.font.size = Pt(12)

doc.add_page_break()

# ── Introduction ──
p = doc.add_paragraph()
run = p.add_run('Introduction')
run.bold = True
run.font.size = Pt(18)
run.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)

intro_text = (
    "Ce rapport présente le bilan des dix séances de Mise en Situations Professionnelles (MSP) "
    "effectuées au Collège Ibn Toumert de Marrakech. Chaque séance a été l'occasion de mettre en pratique "
    "les compétences pédagogiques, didactiques et relationnelles acquises au CRMEF Ibn Roch.\n\n"
    "Les fiches qui suivent adoptent une structure uniforme : contexte d'intervention, déroulement, "
    "analyse selon trois axes (pédagogique, didactique, relationnel) et bilan personnel. "
    "L'ensemble retrace ma progression tout au long de l'année, depuis les premières prises en charge "
    "jusqu'à une maîtrise plus affirmée de la classe.\n\n"
    "Le stage s'est déroulé chaque jeudi, en rotation avec trois autres stagiaires, "
    "auprès d'une classe de 2ème année collège."
)
p = doc.add_paragraph(intro_text)
p.paragraph_format.first_line_indent = Cm(1)
p.paragraph_format.line_spacing = Pt(16)

doc.add_page_break()

# ── Présentation établissement ──
p = doc.add_paragraph()
run = p.add_run('Présentation de l\'établissement')
run.bold = True
run.font.size = Pt(18)
run.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)

pres = (
    "Le Collège Ibn Toumert est situé au centre de Marrakech. Il accueille environ 800 élèves "
    "répartis sur les trois niveaux du cycle collégial. L'établissement dispose de salles "
    "multimédias, d'un laboratoire de sciences et d'une salle d'informatique équipée de 15 postes.\n\n"
    "La classe attribuée pour le stage est une 2ème année collège (2AC) d'environ 35 élèves, "
    "avec un niveau hétérogène. La discipline est l'Informatique, matière évaluée en contrôle continu."
)
p = doc.add_paragraph(pres)
p.paragraph_format.first_line_indent = Cm(1)
p.paragraph_format.line_spacing = Pt(16)

doc.add_page_break()

# ── Fiches de séance ──
for n in range(1, 11):
    is_example = n == 1

    p = doc.add_paragraph()
    run = p.add_run(f'Séance N°{n}')
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)

    # ── Info box ──
    table_info = doc.add_table(rows=4, cols=4)
    table_info.style = 'Table Grid'
    table_info.alignment = WD_TABLE_ALIGNMENT.CENTER
    info_data = [
        ['Date', is_example and 'Jeudi 10 octobre 2025' or '[...]', 'Horaire', is_example and '09h00 — 10h00' or '[...]'],
        ['Matière', is_example and 'Informatique' or '[...]', 'Durée', is_example and '50 min' or '[...]'],
        ['Niveau', '2AC', 'Effectif', is_example and '34 élèves' or '[...]'],
        ['Thème', is_example and 'Introduction aux réseaux informatiques' or '[...]', 'Type', is_example and 'Cours magistral + exercice' or '[...]'],
    ]
    for i, row_data in enumerate(info_data):
        for j, val in enumerate(row_data):
            cell = table_info.rows[i].cells[j]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(val)
            run.font.size = Pt(10)
            if j % 2 == 0:
                run.bold = True
                set_cell_shading(cell, 'E8EDF2')

    doc.add_paragraph('')

    # ── 1. Contexte ──
    p = doc.add_paragraph()
    run = p.add_run('1. Contexte')
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)

    if is_example:
        ctx = (
            "Cette séance est la première que j'ai prise en charge seul. Elle s'inscrit dans le cadre "
            "du module Réseaux Informatiques (S2). Avant cette séance, les élèves avaient déjà abordé "
            "la notion de réseau local en cours avec le tuteur. L'objectif était de consolider ces "
            "connaissances et d'introduire les notions de protocole et d'adressage IP.\n\n"
            "J'avais préparé un diaporama simple et un exercice de reconnaissance de schéma de réseau. "
            "J'appréhendais surtout la gestion du temps et la capacité à capter l'attention de la classe."
        )
    else:
        ctx = (
            "[Décrire ici le contexte de la séance : où en était la progression ? Quel objectif ? "
            "Quelles étaient les difficultés anticipées ? Quelle préparation ?]"
        )
    p = doc.add_paragraph(ctx)
    p.paragraph_format.first_line_indent = Cm(0.5)
    p.paragraph_format.space_after = Pt(6)

    # ── 2. Déroulement ──
    p = doc.add_paragraph()
    run = p.add_run('2. Déroulement de la séance')
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)

    if is_example:
        t_data = [['Phase', 'Durée', 'Activité', 'Observations'],
                  ['Accroche', '5 min', 'Question : « Qu\'est-ce qu\'Internet ? »', 'Plusieurs réponses : « Google », « le wifi ». Noté les représentations initiales.'],
                  ['Rappel', '5 min', 'Reprise du cours précédent (définition réseau)', 'Élèves participent bien, 3 volontaires au tableau pour dessiner un réseau.'],
                  ['Nouveau', '20 min', 'Cours dialogué : protocole, adresse IP', 'Bon accrochage sur l\'analogie de la poste. Quelques confusions IP/MAC. J\'ai reformulé avec un schéma.'],
                  ['Exercice', '10 min', 'Reconnaître un type de réseau sur schéma', 'Exercice fait en binôme. Je circule, aide les groupes faibles. Taux de réussite ~70%.'],
                  ['Synthèse', '10 min', 'Correction collective + trace écrite', 'Distribution de la trace écrite pré-imprimée pour gagner du temps.'],
        ]
        table = doc.add_table(rows=len(t_data), cols=4)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        for i, row_data in enumerate(t_data):
            for j, val in enumerate(row_data):
                cell = table.rows[i].cells[j]
                cell.text = ''
                p = cell.paragraphs[0]
                run = p.add_run(val)
                run.font.size = Pt(10)
                if i == 0:
                    run.bold = True
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    set_cell_shading(cell, '1E3A5F')
    else:
        p = doc.add_paragraph('[Insérer ici un tableau de déroulement : Phase | Durée | Activité | Observations]')

    doc.add_paragraph('')

    # ── 3. Analyse par axe ──
    p = doc.add_paragraph()
    run = p.add_run('3. Analyse par axe')
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)

    axes = [
        ('Axe pédagogique', is_example and (
            "Le temps a été globalement respecté, bien que la phase de synthèse ait pris 5 minutes "
            "de plus que prévu car la correction a soulevé des questions. J'ai utilisé le tableau "
            "pour le schéma et le vidéoprojecteur pour le diaporama. La circulation entre les rangs "
            "pendant l'exercice m'a permis de repérer les élèves en difficulté. À améliorer : "
            "distribuer la trace écrite en début d'exercice plutôt qu'à la fin, pour permettre "
            "aux élèves de s'y référer pendant l'exercice."
        ) or '[Analyser ici : gestion du temps, supports, organisation, imprévus]'),
        ('Axe didactique', is_example and (
            "Le concept de protocole était nouveau et abstrait. L'analogie postale a bien fonctionné "
            "pour la majorité. Environ 40% des élèves confondaient encore adresse IP et adresse MAC "
            "à l'exercice. J'ai dû repasser par le schéma pour clarifier. La trace écrite pré-imprimée "
            "était complète mais trop dense — mieux vaut laisser des espaces à remplir pour maintenir "
            "l'attention pendant la synthèse."
        ) or '[Analyser ici : concepts visés, difficultés, stratégies, résultats]'),
        ('Axe relationnel', is_example and (
            "Le climat était positif. J'ai senti une certaine nervosité au début, mais elle s'est "
            "estompée après les 5 premières minutes. Le tuteur m'a laissé gérer la classe et n'est "
            "intervenu qu'en fin de séance pour féliciter les élèves. J'ai veillé à varier les "
            "sollicitations (volontaires, désignés, binômes). Un élève au fond perturbait : je me "
            "suis approché de son bureau sans le réprimander publiquement — cela a suffi."
        ) or '[Analyser ici : posture, discipline, interactions, climat, feedback tuteur]'),
    ]
    for title, content in axes:
        p = doc.add_paragraph()
        run = p.add_run(f'{title}')
        run.bold = True
        run.font.size = Pt(12)
        p = doc.add_paragraph(content)
        p.paragraph_format.first_line_indent = Cm(0.5)
        p.paragraph_format.space_after = Pt(8)
        p.paragraph_format.line_spacing = Pt(14)

    # ── 4. Bilan ──
    p = doc.add_paragraph()
    run = p.add_run('4. Bilan et pistes d\'amélioration')
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)

    if is_example:
        bilan = (
            "Première séance encourageante. J'ai validé ma capacité à tenir une classe et à structurer "
            "une séance. Les principaux axes d'amélioration : alléger la trace écrite, prévoir un "
            "exercice différencié pour les plus rapides, et chronométrer plus rigoureusement chaque phase."
        )
    else:
        bilan = "[Bilan : ce que j'ai appris, ce que je dois améliorer, objectif pour la prochaine séance]"
    p = doc.add_paragraph(bilan)
    p.paragraph_format.first_line_indent = Cm(0.5)
    p.paragraph_format.space_after = Pt(12)

    if n < 10:
        doc.add_page_break()

# ── Synthèse générale ──
doc.add_page_break()
p = doc.add_paragraph()
run = p.add_run('Synthèse et bilan général')
run.bold = True
run.font.size = Pt(18)
run.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)

synth = (
    "Les dix séances de MSP m'ont permis de passer d'une appréhension légitime à une posture "
    "enseignante plus assurée. J'ai progressé sur plusieurs plans :\n\n"
    "• Pédagogique : meilleure gestion du temps, diversification des supports, anticipation des transitions.\n"
    "• Didactique : capacité à reformuler, à adapter mon discours au niveau des élèves, à concevoir des exercices progressifs.\n"
    "• Relationnel : maintien d'un climat de classe serein, gestion discrète des perturbations, communication bienveillante.\n\n"
    "Je retiens que l'alternance des activités courtes (max 15 minutes par phase) est la clé pour maintenir "
    "l'attention des collégiens. L'observation du tuteur et ses retours réguliers ont été déterminants "
    "dans ma progression."
)
p = doc.add_paragraph(synth)
p.paragraph_format.first_line_indent = Cm(1)
p.paragraph_format.line_spacing = Pt(16)

# ── Save ──
output_path = os.path.join(os.path.dirname(__file__), 'Rapport_MSP_Exemple.docx')
doc.save(output_path)
print(f'Document généré : {output_path}')
