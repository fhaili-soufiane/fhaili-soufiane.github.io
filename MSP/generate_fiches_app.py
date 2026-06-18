from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

doc = Document()

style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

def set_cell_shading(cell, color):
    shading = OxmlElement('w:shd')
    shading.set(qn('w:val'), 'clear')
    shading.set(qn('w:color'), 'auto')
    shading.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading)

def add_table(doc, rows_data, header, col_widths=None):
    table = doc.add_table(rows=len(rows_data) + 1, cols=len(header))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(header):
        cell = table.rows[0].cells[j]
        cell.text = ''
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(10)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, '1E3A5F')
    for i, row_data in enumerate(rows_data):
        for j, val in enumerate(row_data):
            cell = table.rows[i + 1].cells[j]
            cell.text = ''
            run = cell.paragraphs[0].add_run(val)
            run.font.size = Pt(10)
    if col_widths:
        for row in table.rows:
            for j, w in enumerate(col_widths):
                row.cells[j].width = Cm(w)
    doc.add_paragraph('')
    return table

def add_heading(doc, text, level=13):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(level)
    run.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)

def add_body(doc, text, indent=True):
    p = doc.add_paragraph(text)
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.5)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = Pt(14)

# ═══════════════════════════════════════
# COVER PAGE
# ═══════════════════════════════════════
for _ in range(6):
    doc.add_paragraph('')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Analyse des Pratiques Professionnelles\nMises en Situations Professionnelles')
run.bold = True
run.font.size = Pt(24)
run.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)

doc.add_paragraph('')
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('CRMEF Ibn Roch Marrakech\nFilière Collège — Spécialité Informatique\nAnnée de Formation 2025-2026')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

doc.add_paragraph('')
doc.add_paragraph('')

for label, value in [
    ('Stagiaire :', 'Soufiane FHAILI'),
    ('Établissement :', 'Collège Ibn Toumert, Marrakech'),
    ('Niveau :', '2ème année collège'),
    ('Tuteur :', 'M. El Amrani'),
    ('Période :', 'Octobre 2025 — Juin 2026'),
]:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f'{label} ')
    r.bold = True
    r.font.size = Pt(12)
    r = p.add_run(value)
    r.font.size = Pt(12)

doc.add_page_break()

# ═══════════════════════════════════════
# INTRODUCTION
# ═══════════════════════════════════════
add_heading(doc, 'Introduction', 18)
add_body(doc, (
    "Ce document regroupe les dix fiches d'analyse des pratiques professionnelles réalisées "
    "dans le cadre des Mises en Situations Professionnelles (MSP) au Collège Ibn Toumert de Marrakech. "
    "Chaque fiche s'appuie sur la grille d'observation officielle du module APP et mobilise "
    "les trois dimensions : didactique, pédagogique et relationnelle.\n\n"
    "La démarche adoptée suit les étapes de l'analyse réflexive : description objective, "
    "problématisation, analyse multidimensionnelle, théorisation et réinvestissement."
))

doc.add_page_break()

# ═══════════════════════════════════════
# PRÉSENTATION ÉTABLISSEMENT
# ═══════════════════════════════════════
add_heading(doc, "Présentation de l'établissement", 18)
add_body(doc, (
    "Le Collège Ibn Toumert est situé au centre de Marrakech. Il accueille environ 800 élèves "
    "répartis sur les trois niveaux du cycle collégial. L'établissement dispose d'une salle "
    "d'informatique équipée de 15 postes, d'un vidéoprojecteur mobile et d'une connexion Internet.\n\n"
    "La classe attribuée pour le stage est une 2ème année collège (2AC) d'environ 35 élèves, "
    "avec un niveau hétérogène. La discipline est l'Informatique, évaluée en contrôle continu. "
    "Les séances ont lieu chaque jeudi, en rotation avec trois autres stagiaires."
))

doc.add_page_break()

# ═══════════════════════════════════════
# FICHES (10)
# ═══════════════════════════════════════
for n in range(1, 11):
    ex = n == 1  # True only for session 1 (example)

    add_heading(doc, f"Fiche d'analyse N°{n}", 18)
    p = doc.add_paragraph()
    r = p.add_run(ex and 'Séance observée — Exemple' or '[Type : Observation / Co-animation / Prise en charge]')
    r.italic = True
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    doc.add_paragraph('')

    # ── 1. Informations générales ──
    add_heading(doc, '1. Informations générales', 13)

    info_data = ex and [
        ['Stagiaire', 'Soufiane FHAILI', 'Établissement', 'Collège Ibn Toumert'],
        ['Classe', '2AC', 'Effectif', '34 élèves'],
        ['Date', 'Jeudi 10 octobre 2025', 'Horaire', '09h00 — 10h00'],
        ['Thème', 'Introduction aux réseaux informatiques', 'Type', 'Observation'],
        ['Durée', '50 min', 'Tuteur', 'M. El Amrani'],
    ] or [
        ['Stagiaire', 'Soufiane FHAILI', 'Établissement', 'Collège Ibn Toumert'],
        ['Classe', '2AC', 'Effectif', '[...]'],
        ['Date', '[...]', 'Horaire', '[...]'],
        ['Thème', '[...]', 'Type', '[...]'],
        ['Durée', '[...]', 'Tuteur', 'M. El Amrani'],
    ]
    table = doc.add_table(rows=5, cols=4)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, row_data in enumerate(info_data):
        for j, val in enumerate(row_data):
            cell = table.rows[i].cells[j]
            cell.text = ''
            run = cell.paragraphs[0].add_run(val)
            run.font.size = Pt(10)
            if j % 2 == 0:
                run.bold = True
                set_cell_shading(cell, 'E8EDF2')
    doc.add_paragraph('')

    # ── 2. Contexte ──
    add_heading(doc, '2. Contexte de la séance (Description)', 13)
    add_body(doc, ex and (
        "Cette séance constitue la première MSP. Elle s'inscrit dans le module Réseaux "
        "Informatiques (S2). Avant cette séance, les élèves avaient abordé la notion de réseau "
        "local avec le tuteur. L'objectif est de consolider ces connaissances et d'introduire "
        "les notions de protocole et d'adressage IP.\n\n"
        "Le tuteur a pris en charge l'intégralité de la séance. Mon rôle s'est limité à observer "
        "et à noter le déroulement, les interactions et les stratégies employées."
    ) or (
        "[Décrire objectivement : où en est la progression ? Quel est l'objectif de la séance ? "
        "Quel est mon rôle : observer, co-animer ou prendre en charge ?]"
    ))

    # ── 3. Déroulement ──
    add_heading(doc, '3. Déroulement chronologique', 13)
    if ex:
        add_table(doc, [
            ['Accroche', '5 min', 'Question orale : « Qu\'est-ce qu\'Internet ? »',
             'Réponses : « Google », « le wifi »', 'Le tuteur note les représentations au tableau'],
            ['Rappel', '5 min', 'Reprise du cours précédent', '3 volontaires au tableau',
             'Bonne participation, le tuteur corrige en douceur'],
            ['Nouveau', '20 min', 'Cours dialogué : protocole, adresse IP',
             'Écoute, prise de notes, questions', 'Analogie postale efficace. Confusions IP/MAC'],
            ['Exercice', '10 min', 'Schéma à compléter', 'Travail en binômes',
             'Taux de réussite ~70%. Le tuteur circule et aide'],
            ['Synthèse', '10 min', 'Correction collective + trace écrite', 'Copie de la trace',
             'Trace pré-imprimée : gain de temps'],
        ], ['Phase', 'Durée', 'Activité', 'Activité élèves', 'Observations'])
    else:
        add_body(doc, '[Insérer le tableau de déroulement : Phase | Durée | Activité | Activité élèves | Observations]')

    # ── 4. Problématisation ──
    add_heading(doc, '4. Problématisation', 13)
    add_body(doc, ex and (
        "Plusieurs questions émergent de cette observation :\n"
        "• Comment rendre accessible un concept abstrait (protocole) à des élèves de 2AC ?\n"
        "• Quel est l'impact de l'analogie sur la compréhension à long terme ?\n"
        "• Comment gérer l'hétérogénéité des niveaux au sein d'une même classe ?\n"
        "• La trace écrite pré-imprimée favorise-t-elle vraiment l'apprentissage ?"
    ) or (
        "[Identifier les questions et problèmes soulevés par l'observation]"
    ))

    # ── 5. Analyse par dimensions ──
    add_heading(doc, '5. Analyse par dimensions', 13)
    if ex:
        add_table(doc, [
            ['Didactique',
             'Situation de départ, construction des apprentissages, gestion des erreurs, évaluation',
             'Situation de départ : question orale pour faire émerger les représentations. '
             'Construction : progression du concret (réseau local) vers l\'abstrait (protocole). '
             'Gestion des erreurs : confusion IP/MAC traitée par retour au schéma. '
             'Évaluation : exercice en binôme, taux de réussite ~70%.',
             'Prévoir un exercice différencié. Ajouter des espaces à remplir dans la trace écrite.'],
            ['Pédagogique',
             'Gestion du temps, méthodes pédagogiques, diversification des activités',
             'Temps globalement respecté (50 min). Alternance de 5 phases (question, rappel, cours, '
             'exercice, synthèse). Méthode : cours dialogué + travail en binôme. '
             'Supports : tableau, vidéoprojecteur, fiche imprimée.',
             'Distribuer la fiche plus tôt. Chronométrer chaque phase.'],
            ['Relationnel',
             'Communication, motivation, gestion de classe, interaction',
             'Climat serein. Posture calme et autoritaire. Variété des sollicitations '
             '(volontaires + désignés). Élève perturbateur géré par rapprochement discret. '
             'Le stagiaire a été inclus en fin de séance.',
             'Impliquer davantage les élèves en difficulté.'],
        ], ['Dimension', 'Éléments observés', 'Observations', "Pistes d'amélioration"],
           col_widths=[2.5, 4, 5.5, 4.5])
    else:
        add_table(doc, [
            ['Didactique',
             'Situation de départ, construction des apprentissages, gestion des erreurs, évaluation',
             '[...]', '[...]'],
            ['Pédagogique',
             'Gestion du temps, méthodes pédagogiques, diversification des activités',
             '[...]', '[...]'],
            ['Relationnel',
             'Communication, motivation, gestion de classe, interaction',
             '[...]', '[...]'],
        ], ['Dimension', 'Éléments observés', 'Observations', "Pistes d'amélioration"],
           col_widths=[2.5, 4, 5.5, 4.5])

    # ── 6. Théorisation ──
    add_heading(doc, '6. Théorisation', 13)
    add_body(doc, ex and (
        "Plusieurs concepts pédagogiques éclairent cette séance :\n"
        "• **Conflit cognitif** (Piaget) : la confusion IP/MAC a créé un déséquilibre, "
        "amenant les élèves à restructurer leurs connaissances.\n"
        "• **Médiation** (Vygotsky) : le tuteur a joué un rôle de médiateur en utilisant "
        "l'analogie postale pour faire le lien entre connu et inconnu.\n"
        "• **Pédagogie différenciée** : l'exercice en binôme permet une entraide, "
        "mais un prolongement pour les plus rapides manquait.\n"
        "• **Évaluation formative** : la correction collective et la circulation ont permis "
        "un feedback immédiat."
    ) or (
        "[Mobiliser des références théoriques : concepts, auteurs, modèles pédagogiques "
        "qui permettent d'interpréter les faits observés]"
    ))

    # ── 7. Synthèse réflexive ──
    add_heading(doc, '7. Synthèse réflexive personnelle', 13)
    add_body(doc, ex and (
        "Cette première observation m'a permis de comprendre l'importance de structurer "
        "la séance en phases courtes et variées. L'analogie comme outil didactique est "
        "particulièrement efficace pour les concepts abstraits. Je retiens aussi que "
        "la gestion discrète des perturbations est plus efficace qu'une réprimande publique.\n\n"
        "Pour ma prochaine prise en charge, je veillerai à : préparer une trace écrite "
        "interactive avec des espaces à remplir, prévoir une activité de prolongement "
        "pour les plus rapides, et chronométrer chaque phase."
    ) or (
        "[Rédiger une synthèse réflexive : ce que j'ai appris, ce que je remets en question, "
        "ce que je réinvestirai dans ma pratique]"
    ))

    if n < 10:
        doc.add_page_break()

# ═══════════════════════════════════════
# SYNTHÈSE GÉNÉRALE
# ═══════════════════════════════════════
doc.add_page_break()
add_heading(doc, 'Synthèse générale', 18)
add_body(doc, (
    "Les dix séances de MSP ont permis de développer une posture réflexive structurée "
    "autour des trois dimensions de l'analyse des pratiques professionnelles :\n\n"
    "**Dimension didactique :** progression dans la capacité à concevoir des situations "
    "d'apprentissage adaptées, à gérer les erreurs comme leviers et à diversifier les "
    "modalités d'évaluation.\n\n"
    "**Dimension pédagogique :** amélioration de la gestion du temps, diversification "
    "des méthodes (cours dialogué, travail en groupe, exercices différenciés) et "
    "utilisation raisonnée des supports.\n\n"
    "**Dimension relationnelle :** construction d'une posture enseignante équilibrée, "
    "gestion discrète des perturbations, communication bienveillante et exigeante.\n\n"
    "L'alternance des phases courtes (max 15 min) et la variété des activités sont "
    "les clés pour capter l'attention des collégiens. La théorisation des pratiques "
    "(conflit cognitif, médiation, évaluation formative) a permis de donner du sens "
    "aux observations et d'orienter les améliorations."
))

# ── Save ──
output_path = '/tmp/Fiches_APP.docx'
doc.save(output_path)
print(f'Document généré : {output_path}')
