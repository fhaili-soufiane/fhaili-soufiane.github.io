from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

doc = Document()

# ── Styles ──
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

# ── Helpers ──
def set_cell_shading(cell, color):
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:val'), 'clear')
    shading_elm.set(qn('w:color'), 'auto')
    shading_elm.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading_elm)

def add_header_row(table, texts, color='1E3A5F'):
    row = table.rows[0]
    for i, text in enumerate(texts):
        cell = row.cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(text)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(10)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, color)

def add_row(table, cells_texts):
    row = table.add_row()
    for i, text in enumerate(cells_texts):
        cell = row.cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(text)
        run.font.size = Pt(10)

# ── Cover page ──
for _ in range(6):
    doc.add_paragraph('')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Fiches d\'Observation MSP\nMise en Situations Professionnelles')
run.bold = True
run.font.size = Pt(26)
run.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)

doc.add_paragraph('')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('CRMEF Ibn Roch Marrakech\nFilière Collège — Spécialité Informatique\nAnnée de Formation 2025-2026')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

doc.add_paragraph('')
doc.add_paragraph('')

info = [
    ('Stagiaire :', 'Soufiane FHAILI'),
    ('Établissement :', 'Collège Ibn Toumert, Marrakech'),
    ('Niveau :', '2ème année collège'),
    ('Tuteur :', 'M. El Amrani'),
    ('Période :', 'Octobre 2025 — Juin 2026'),
    ('Soutenance :', '19 Juillet 2026'),
]
for label, value in info:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f'{label} ')
    run.bold = True
    run.font.size = Pt(12)
    run = p.add_run(value)
    run.font.size = Pt(12)

doc.add_page_break()

# ── Introduction ──
p = doc.add_paragraph()
run = p.add_run('Introduction')
run.bold = True
run.font.size = Pt(18)
run.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)

intro = (
    "Ce document regroupe les dix fiches d'observation réalisées dans le cadre des Mises en Situations "
    "Professionnelles (MSP) au Collège Ibn Toumert de Marrakech. Chaque fiche rend compte de manière "
    "objective du déroulement d'une séance d'informatique en classe de 2ème année collège.\n\n"
    "Les séances ont alterné entre observations du tuteur, co-animation et prise en charge complète. "
    "Chaque fiche suit une structure uniforme : informations générales, contexte, déroulement chronologique "
    "(tableau), analyse selon trois axes (pédagogique, didactique, relationnel) et bilan.\n\n"
    "L'objectif est de documenter la progression tout au long de l'année et de développer une posture "
    "réflexive sur la pratique enseignante."
)
p = doc.add_paragraph(intro)
p.paragraph_format.first_line_indent = Cm(1)
p.paragraph_format.line_spacing = Pt(16)

doc.add_page_break()

# ── Présentation établissement ──
p = doc.add_paragraph()
run = p.add_run('Présentation de l\'établissement')
run.bold = True
run.font.size = Pt(18)
run.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)

pres = (
    "Le Collège Ibn Toumert est situé au centre de Marrakech. Il accueille environ 800 élèves "
    "répartis sur les trois niveaux du cycle collégial. L'établissement dispose d'une salle "
    "d'informatique équipée de 15 postes, d'un vidéoprojecteur mobile et d'une connexion Internet.\n\n"
    "La classe attribuée pour le stage est une 2ème année collège (2AC) d'environ 35 élèves, "
    "avec un niveau hétérogène. La discipline est l'Informatique, évaluée en contrôle continu. "
    "Les séances ont lieu chaque jeudi, en rotation avec trois autres stagiaires."
)
p = doc.add_paragraph(pres)
p.paragraph_format.first_line_indent = Cm(1)
p.paragraph_format.line_spacing = Pt(16)

doc.add_page_break()

# ── Fiches de séance ──
for n in range(1, 11):
    is_example = n == 1

    p = doc.add_paragraph()
    run = p.add_run(f'Fiche d\'observation N°{n}')
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)

    # Description line
    p = doc.add_paragraph()
    run = p.add_run(is_example and 'Séance observée — Exemple' or '[Type : Observation / Co-animation / Prise en charge]')
    run.italic = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_paragraph('')

    # ── 1. Informations générales ──
    p = doc.add_paragraph()
    run = p.add_run('1. Informations générales')
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)

    table_info = doc.add_table(rows=5, cols=4)
    table_info.style = 'Table Grid'
    table_info.alignment = WD_TABLE_ALIGNMENT.CENTER

    if is_example:
        info_data = [
            ['Stagiaire', 'Soufiane FHAILI', 'Établissement', 'Collège Ibn Toumert'],
            ['Classe', '2AC', 'Effectif', '34 élèves'],
            ['Date', 'Jeudi 10 octobre 2025', 'Horaire', '09h00 — 10h00'],
            ['Thème', 'Introduction aux réseaux', 'Type', 'Observation'],
            ['Durée', '50 min', 'Tuteur', 'M. El Amrani'],
        ]
    else:
        info_data = [
            ['Stagiaire', 'Soufiane FHAILI', 'Établissement', 'Collège Ibn Toumert'],
            ['Classe', '2AC', 'Effectif', '[...]'],
            ['Date', '[...]', 'Horaire', '[...]'],
            ['Thème', '[...]', 'Type', '[...]'],
            ['Durée', '[...]', 'Tuteur', 'M. El Amrani'],
        ]

    for i, row_data in enumerate(info_data):
        for j, val in enumerate(row_data):
            cell = table_info.rows[i].cells[j]
            cell.text = ''
            pr = cell.paragraphs[0]
            run = pr.add_run(val)
            run.font.size = Pt(10)
            if j % 2 == 0:
                run.bold = True
                set_cell_shading(cell, 'E8EDF2')

    doc.add_paragraph('')

    # ── 2. Contexte ──
    p = doc.add_paragraph()
    run = p.add_run('2. Contexte de la séance')
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)

    if is_example:
        ctx = (
            "Cette séance constitue la première MSP. Elle s'inscrit dans le cadre du module Réseaux "
            "Informatiques (S2). Avant cette séance, les élèves avaient abordé la notion de réseau "
            "local avec le tuteur. L'objectif de cette séance est de consolider ces connaissances "
            "et d'introduire les notions de protocole et d'adressage IP.\n\n"
            "Le tuteur a pris en charge l'intégralité de la séance. Mon rôle s'est limité à observer "
            "et à noter le déroulement, les interactions et les stratégies employées."
        )
    else:
        ctx = (
            "[Décrire ici le contexte : où en est la progression ? Quel est l'objectif de la séance ? "
            "Quel est mon rôle : observer, co-animer ou prendre en charge ? Y a-t-il des difficultés "
            "anticipées ?]"
        )
    p = doc.add_paragraph(ctx)
    p.paragraph_format.first_line_indent = Cm(0.5)
    p.paragraph_format.space_after = Pt(6)

    # ── 3. Déroulement ──
    p = doc.add_paragraph()
    run = p.add_run('3. Déroulement chronologique')
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)

    if is_example:
        t_data = [
            ['Phase', 'Durée', 'Activité enseignant', 'Activité élèves', 'Observations'],
            ['Accroche', '5 min', 'Question orale : « Qu\'est-ce qu\'Internet ? »', 'Réponses variées : « Google », « le wifi »', 'Le tuteur note les représentations initiales au tableau'],
            ['Rappel', '5 min', 'Reprise du cours précédent', '3 volontaires dessinent un réseau au tableau', 'Bonne participation, le tuteur corrige en douceur'],
            ['Nouveau', '20 min', 'Cours dialogué : protocole, adresse IP', 'Écoute, prise de notes, questions', 'Analogie de la poste très efficace. Certains confondent IP/MAC'],
            ['Exercice', '10 min', 'Distribution d\'un schéma à compléter', 'Travail en binômes', 'Taux de réussite ~70%. Le tuteur circule et aide'],
            ['Synthèse', '10 min', 'Correction collective + trace écrite', 'Copie de la trace', 'Distribution de la trace pré-imprimée : gain de temps'],
        ]
        table = doc.add_table(rows=len(t_data), cols=5)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        for i, row_data in enumerate(t_data):
            for j, val in enumerate(row_data):
                cell = table.rows[i].cells[j]
                cell.text = ''
                pr = cell.paragraphs[0]
                run = pr.add_run(val)
                run.font.size = Pt(10)
                if i == 0:
                    run.bold = True
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                    pr.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    set_cell_shading(cell, '1E3A5F')
    else:
        p = doc.add_paragraph('[Insérer ici le tableau de déroulement : Phase | Durée | Activité enseignant | Activité élèves | Observations]')

    doc.add_paragraph('')

    # ── 4. Analyse par axe ──
    p = doc.add_paragraph()
    run = p.add_run('4. Analyse par axe')
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)

    axes_data = [
        ('Pédagogique', is_example and (
            "Le tuteur a respecté le timing prévu (50 min). L'alternance des activités (question orale, "
            "cours dialogué, exercice en binôme, correction collective) a maintenu l'attention. "
            "Tableau + vidéoprojecteur utilisés. Circulation pendant l'exercice pour suivi individualisé. "
            "À améliorer : distribuer la trace écrite plus tôt."
        ) or '[Organisation, supports, gestion du temps, transitions, imprévus]'),
        ('Didactique', is_example and (
            "Concept de protocole abstrait. Analogie postale efficace pour 60%. 40% confondaient "
            "IP/MAC → retour au schéma. Trace écrite dense — prévoir des espaces à remplir. "
            "Exercice adapté mais pas de prolongement pour les plus rapides."
        ) or '[Concepts visés, stratégies, difficultés, remédiation, résultats]'),
        ('Relationnel', is_example and (
            "Climat serein. Posture calme et autoritaire. Sollicitation variée (volontaires + désignés). "
            "Élève perturbateur : approche discrète suffisante. Le tuteur m'a inclus en cours de séance."
        ) or '[Climat, posture, discipline, communication, interactions]'),
    ]
    axes_table = doc.add_table(rows=len(axes_data) + 1, cols=2)
    axes_table.style = 'Table Grid'
    axes_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # header
    set_cell_shading(axes_table.rows[0].cells[0], '1E3A5F')
    set_cell_shading(axes_table.rows[0].cells[1], '1E3A5F')
    for i, h in enumerate(['Axe', 'Observations']):
        cell = axes_table.rows[0].cells[i]
        cell.text = ''
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(10)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for i, (title, content) in enumerate(axes_data):
        row = axes_table.rows[i + 1]
        # title cell
        cell0 = row.cells[0]
        cell0.text = ''
        run = cell0.paragraphs[0].add_run(title)
        run.bold = True
        run.font.size = Pt(10)
        set_cell_shading(cell0, 'E8EDF2')
        # content cell
        cell1 = row.cells[1]
        cell1.text = ''
        run = cell1.paragraphs[0].add_run(content)
        run.font.size = Pt(10)
    # Set column widths
    for row in axes_table.rows:
        row.cells[0].width = Cm(3.5)
        row.cells[1].width = Cm(13)
    doc.add_paragraph('')

    # ── 5. Bilan ──
    p = doc.add_paragraph()
    run = p.add_run('5. Bilan')
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)

    if is_example:
        bilan = (
            "Points positifs : bonne gestion du temps, alternance efficace des activités, "
            "analogie pertinente pour un concept abstrait.\n"
            "Points à améliorer : prévoir un exercice différencié pour les élèves rapides, "
            "alléger la trace écrite ou la rendre interactive.\n"
            "Objectif pour la prochaine séance : préparer un support avec des espaces à remplir."
        )
    else:
        bilan = (
            "Points positifs : [...]\n"
            "Points à améliorer : [...]\n"
            "Objectif pour la prochaine séance : [...]"
        )
    p = doc.add_paragraph(bilan)
    p.paragraph_format.first_line_indent = Cm(0.5)
    p.paragraph_format.space_after = Pt(12)

    if n < 10:
        doc.add_page_break()

# ── Synthèse générale ──
doc.add_page_break()
p = doc.add_paragraph()
run = p.add_run('Synthèse générale')
run.bold = True
run.font.size = Pt(18)
run.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)

synth = (
    "Les dix séances de MSP ont permis d'observer une progression nette dans la pratique "
    "enseignante, tant chez le tuteur que dans ma propre posture :\n\n"
    "Pédagogique : maîtrise progressive de la gestion du temps, diversification des supports, "
    "anticipation des transitions.\n"
    "Didactique : capacité croissante à reformuler, adapter le discours, concevoir des exercices "
    "progressifs et différenciés.\n"
    "Relationnel : maintien d'un climat de classe favorable, gestion discrète des perturbations, "
    "communication bienveillante mais ferme.\n\n"
    "L'alternance des phases courtes (max 15 min) et la variété des activités sont les clés "
    "pour capter l'attention des collégiens. Les retours réguliers du tuteur ont été déterminants "
    "dans la progression."
)
p = doc.add_paragraph(synth)
p.paragraph_format.first_line_indent = Cm(1)
p.paragraph_format.line_spacing = Pt(16)

# ── Save ──
output_path = '/tmp/Fiches_Observation_MSP.docx'
doc.save(output_path)
print(f'Document généré : {output_path}')
